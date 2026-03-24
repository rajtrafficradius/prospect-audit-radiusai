#!/usr/bin/env python3
"""
Prospect Audit & Strategy — Branded DOCX Generator (SEO + AEO + GEO)
=====================================================================
Generates a 15-25 page branded strategy document with 12 sections
covering all three layers of modern search.

Adapt this template for each prospect by replacing placeholder data
with real data from the JSON/CSV files generated in Phases 1-5.

Usage: python3 create_strategy_docx.py
Requires: python-docx (pip install python-docx)
"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ══════════════════════════════════════════════════════════════
# CONFIGURATION — Update these for each prospect
# ══════════════════════════════════════════════════════════════

import datetime
PROSPECT_NAME = os.environ.get("PROSPECT_NAME", "TrafficRadius Prospect")
PROSPECT_DOMAIN = os.environ.get("PROSPECT_DOMAIN", "example.com")
PROSPECT_DATE = os.environ.get("PROSPECT_DATE", datetime.datetime.now().strftime("%B %d, %Y"))
OUTPUT_DIR = os.path.join(os.environ.get("OUTPUT_DIR", "/home/ubuntu/output"), "deliverables")
CHARTS_DIR = os.path.join(os.environ.get("OUTPUT_DIR", "/home/ubuntu/output"), "charts")
LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'src', 'static', 'logo.png')
DATA_DIR = os.environ.get("OUTPUT_DIR", "/home/ubuntu/output")

try:
    with open(os.path.join(DATA_DIR, "strategy_narrative.json"), "r") as f:
        narrative_data = json.load(f)
except Exception:
    narrative_data = {}

try:
    with open(os.path.join(DATA_DIR, "market_intelligence.json"), "r") as f:
        market_data = json.load(f)
except Exception:
    market_data = {}

try:
    with open(os.path.join(DATA_DIR, "audit_findings.json"), "r") as f:
        audit_data = json.load(f)
except Exception:
    audit_data = {}

try:
    with open(os.path.join(DATA_DIR, "cro_assessment.json"), "r") as f:
        cro_data = json.load(f)
except Exception:
    cro_data = {}

try:
    with open(os.path.join(DATA_DIR, "competitor_shadowing.json"), "r") as f:
        shadow_data = json.load(f)
except Exception:
    shadow_data = {}

# ── Brand Colors ──────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2E, 0x50, 0x90)
LIGHT_BLUE = RGBColor(0x4A, 0x90, 0xD9)
DARK_GREY = RGBColor(0x2D, 0x37, 0x48)
MID_GREY = RGBColor(0x4A, 0x55, 0x68)
LIGHT_GREY = RGBColor(0x71, 0x80, 0x96)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Layer accent colors
SEO_BLUE = RGBColor(0x2E, 0x50, 0x90)
AEO_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GEO_PURPLE = RGBColor(0x9B, 0x59, 0xB6)

NAVY_HEX = "1B2A4A"
BLUE_HEX = "2E5090"
AEO_HEX = "E67E22"
GEO_HEX = "9B59B6"
LIGHT_BG = "F7FAFC"
TABLE_HEADER_BG = "1B2A4A"
TABLE_ALT_ROW = "F0F2F5"

# ══════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_branded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")
    return table

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles['Heading 1']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = NAVY
        run.font.size = Pt(24)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.style = doc.styles['Heading 2']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = BLUE
        run.font.size = Pt(18)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        p.style = doc.styles['Heading 3']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(14)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color if color else DARK_GREY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_callout(doc, text, accent_color_hex=BLUE_HEX):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREY
    run.italic = True
    set_cell_shading(cell, LIGHT_BG)
    set_cell_border(cell,
        left={"sz": "24", "color": accent_color_hex, "val": "single"},
        top={"sz": "4", "color": "E2E8F0", "val": "single"},
        bottom={"sz": "4", "color": "E2E8F0", "val": "single"},
        right={"sz": "4", "color": "E2E8F0", "val": "single"}
    )
    doc.add_paragraph("")

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_chart_image(doc, filename, width=6.0):
    path = os.path.join(CHARTS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        doc.add_paragraph("")

def add_page_break(doc):
    doc.add_page_break()

def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "© Copyright 2026 TrafficRadius, All rights reserved.\t\tPage "
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════
# DOCUMENT CREATION
# ══════════════════════════════════════════════════════════════

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GREY
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(6)
pf.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

add_footer(doc.sections[0])

# ── COVER PAGE ───────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run()
if os.path.exists(LOGO_PATH):
    run.add_picture(LOGO_PATH, width=Inches(2.5))

for _ in range(3):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(PROSPECT_NAME)
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Integrated SEO, AEO & GEO")
run.font.size = Pt(28)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Prospect Audit & Growth Strategy")
run.font.size = Pt(22)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("A Data-Driven Blueprint for Multi-Channel Search Dominance")
run.font.size = Pt(14)
run.font.color.rgb = MID_GREY
run.font.name = 'Calibri'
run.italic = True

for _ in range(4):
    doc.add_paragraph("")

add_divider(doc)

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
meta_data = [
    ("Prepared For:", PROSPECT_NAME),
    ("Prepared By:", "TrafficRadius — trafficradius.com.au"),
    ("Date:", PROSPECT_DATE),
    ("Classification:", "Confidential — Client Use Only"),
]
for i, (label, value) in enumerate(meta_data):
    cell_l = meta_table.rows[i].cells[0]
    cell_l.text = ""
    run_l = cell_l.paragraphs[0].add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_l.font.color.rgb = NAVY
    run_l.font.name = 'Calibri'
    cell_l.width = Inches(1.8)
    cell_r = meta_table.rows[i].cells[1]
    cell_r.text = ""
    run_r = cell_r.paragraphs[0].add_run(value)
    run_r.font.size = Pt(11)
    run_r.font.color.rgb = DARK_GREY
    run_r.font.name = 'Calibri'

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("© Copyright 2026 TrafficRadius, All rights reserved.")
run.font.size = Pt(9)
run.font.color.rgb = LIGHT_GREY
run.font.name = 'Calibri'

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. Executive Summary", 1)

add_body(doc,
    narrative_data.get("executive_summary", 
        f"This document presents a comprehensive, multi-layer digital audit and growth strategy for {PROSPECT_NAME}. "
        "Our analysis spans the three pillars of modern search — Traditional SEO, Answer Engine Optimization (AEO), "
        "and Generative Engine Optimization (GEO) — to identify opportunities across every channel where "
        "potential customers are searching for solutions."
    )
)

add_callout(doc,
    "Key Insight: [PROSPECT_NAME] currently captures less than X% of the total addressable "
    "search market in its core verticals, with significant untapped opportunities in answer "
    "engines and AI-powered search platforms."
)

add_body(doc,
    "Our strategy is built on an integrated framework that addresses all three search layers "
    "through a phased, 90-day action plan. This approach ensures sustainable, long-term growth "
    "across traditional search, featured snippets, and AI-generated answers."
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 2: THE THREE LAYERS OF MODERN SEARCH
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "2. The Three Layers of Modern Search", 1)

add_body(doc,
    "The search landscape has fundamentally changed. Today, businesses must be visible not just "
    "in traditional Google results, but across three distinct layers of search. Understanding "
    "these layers is critical to building an effective digital strategy."
)

add_heading_styled(doc, "Layer 1: SEO — Traditional Search Engine Optimization", 2)
add_body(doc,
    "Traditional SEO focuses on ranking in organic search results on Google, Bing, and other "
    "search engines. This remains the foundation of digital visibility and drives the majority "
    "of high-intent commercial traffic."
)

add_heading_styled(doc, "Layer 2: AEO — Answer Engine Optimization", 2)
add_body(doc,
    "AEO focuses on winning Featured Snippets, People Also Ask (PAA) boxes, and voice search "
    "results. These positions capture significant click-through rates and establish the brand "
    "as the definitive answer to common industry questions."
)

add_heading_styled(doc, "Layer 3: GEO — Generative Engine Optimization", 2)
add_body(doc,
    "GEO is the newest frontier, focusing on being cited and referenced by AI-powered search "
    "platforms such as ChatGPT, Google AI Overviews, Perplexity, and Claude. As more users "
    "turn to AI for answers, GEO ensures the brand is part of the conversation."
)

add_chart_image(doc, "three_layer_overview.png")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 3: BUSINESS & MARKET CONTEXT
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "3. Business & Market Context", 1)

add_heading_styled(doc, "Company Profile", 2)
add_body(doc,
    narrative_data.get("company_profile", 
        f"TrafficRadius identified {PROSPECT_NAME} as a key operator within its niche, requiring an enhanced discovery posture across SEO, AEO, and GEO platforms."
    )
)

add_heading_styled(doc, "Digital Maturity Assessment", 2)
add_body(doc,
    narrative_data.get("digital_maturity_assessment", 
        "Currently, the prospect demonstrates a foundational baseline of SEO assets, however, advanced schema, comprehensive backlink integration, and strict technical AI optimization parameters are lacking. Implementation of the outlined strategies will exponentially grow market share."
    )
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 4: CURRENT DIGITAL PERFORMANCE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "4. Current Digital Performance", 1)

add_body(doc,
    f"The following metrics provide a data-driven baseline of {PROSPECT_NAME}'s current "
    "organic search performance, sourced from SEMrush."
)

prospect_data = market_data.get("prospect", {})
add_branded_table(doc,
    headers=["Metric", "Current Value"],
    rows=[
        ["Domain Authority Score", str(prospect_data.get("authority_score", "N/A"))],
        ["Organic Keywords", str(prospect_data.get("organic_keywords", "N/A"))],
        ["Monthly Organic Traffic", str(prospect_data.get("organic_traffic", "N/A"))],
        ["Organic Traffic Value", f"${prospect_data.get('organic_traffic_value', '0')}"],
        ["Backlinks", f"{prospect_data.get('backlinks', 'N/A'):,}" if isinstance(prospect_data.get('backlinks'), int) else prospect_data.get('backlinks', 'N/A')],
        ["Referring Domains", f"{prospect_data.get('referring_domains', 'N/A'):,}" if isinstance(prospect_data.get('referring_domains'), int) else prospect_data.get('referring_domains', 'N/A')],
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 5: COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "5. Competitive Landscape", 1)

add_body(doc,
    narrative_data.get("competitive_landscape_analysis", 
        f"Understanding the competitive landscape is critical to identifying where {PROSPECT_NAME} "
        "can gain the most ground across all three search layers."
    )
)

comp_rows = [
    [PROSPECT_NAME, str(market_data.get("prospect", {}).get("organic_keywords", "N/A")), str(market_data.get("prospect", {}).get("organic_traffic", "N/A")), str(market_data.get("prospect", {}).get("organic_traffic_value", "N/A")), "N/A"]
]
for comp in market_data.get("competitors", [])[:3]:
    comp_rows.append([comp.get("domain", ""), str(comp.get("organic_keywords", "N/A")), str(comp.get("organic_traffic", "N/A")), str(comp.get("organic_traffic_value", "N/A")), "N/A"])

if len(comp_rows) == 1:
    comp_rows.extend([
        ["Competitor 1", "N/A", "N/A", "N/A"],
        ["Competitor 2", "N/A", "N/A", "N/A"],
        ["Competitor 3", "N/A", "N/A", "N/A"],
    ])

add_branded_table(doc,
    headers=["Competitor", "Organic Keywords", "Monthly Traffic", "Traffic Value", "Backlinks"],
    rows=comp_rows
)

add_chart_image(doc, "competitive_landscape.png")

if shadow_data:
    add_heading_styled(doc, "5.1 Value Proposition Gap Analysis", 2)
    add_body(doc, shadow_data.get("overall_gap_summary", "Competitor analysis indicates several missed opportunities."))
    for gap in shadow_data.get("gaps", []):
        add_callout(doc, f"Versus {gap.get('competitor', '')}:\nThey highlight: {gap.get('what_they_have', '')}\nWhy it wins: {gap.get('why_it_works', '')}\nStrategy Shift: {gap.get('how_to_beat_it', '')}")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 6: INTEGRATED AUDIT — SEO, AEO & GEO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "6. Integrated Audit — SEO, AEO & GEO", 1)

add_body(doc,
    "Our integrated audit assessed the website across all three search layers. "
    "The scorecard below provides a snapshot of current readiness for each layer."
)

add_chart_image(doc, "integrated_scorecard.png")

add_heading_styled(doc, "6.1 SEO Audit Findings", 2)
add_body(doc,
    "The following high-impact technical SEO issues were identified during our audit."
)
qw_rows = []
for qw in narrative_data.get("seo_quick_wins", []):
    qw_rows.append(["High", "Technical SEO", qw.get("title", ""), qw.get("description", "")])

if not qw_rows:
    qw_rows = [
        ["Critical", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
        ["High", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
        ["Medium", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
    ]

add_branded_table(doc,
    headers=["Severity", "Area", "Finding", "Recommendation"],
    rows=qw_rows
)

add_heading_styled(doc, "6.2 AEO Audit Findings", 2)
add_callout(doc,
    "AEO focuses on how well the website is structured to win Featured Snippets, "
    "People Also Ask boxes, and voice search results.",
    AEO_HEX
)
aeo_rows = []
for finding in audit_data.get("aeo_findings", []):
    aeo_rows.append([finding.get("area", "AEO"), finding.get("current_status", "N/A"), finding.get("title", ""), finding.get("recommendation", "")])

if not aeo_rows:
    aeo_rows = [["AEO Readiness", "N/A", "No AEO findings tracked.", "Expand schema deployment."]]

add_branded_table(doc,
    headers=["Area", "Status", "Finding", "Recommendation"],
    rows=aeo_rows
)

add_heading_styled(doc, "6.3 GEO Audit Findings", 2)
add_callout(doc,
    "GEO assesses how visible the brand is to AI-powered search engines like "
    "ChatGPT, Perplexity, and Google AI Overviews.",
    GEO_HEX
)
geo_rows = []
for finding in audit_data.get("geo_findings", []):
    geo_rows.append([finding.get("area", "GEO"), finding.get("current_status", "N/A"), finding.get("title", ""), finding.get("recommendation", "")])

if not geo_rows:
    geo_rows = [["GEO Targets", "N/A", "No GEO findings tracked.", "Build Entity Signals."]]

add_branded_table(doc,
    headers=["Area", "Status", "Finding", "Recommendation"],
    rows=geo_rows
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 7: CRO ASSESSMENT
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "7. CRO Assessment", 1)

add_body(doc,
    "Beyond attracting traffic across all search layers, it is essential that the website "
    "effectively converts visitors into leads or customers."
)

cro_rows = []
for finding in cro_data.get("findings", []):
    cro_rows.append([finding.get("area", ""), finding.get("current_status", ""), finding.get("opportunity", "")])

if not cro_rows:
    cro_rows = [
        ["Primary CTA", "Not Verified", "Awaiting Scan"],
        ["Lead Capture Form", "Not Verified", "Awaiting Scan"],
        ["Trust Signals", "Not Verified", "Awaiting Scan"],
        ["Mobile Experience", "Not Verified", "Awaiting Scan"],
    ]

add_branded_table(doc,
    headers=["Area", "Current Status", "Opportunity"],
    rows=cro_rows
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 8: KEYWORD OPPORTUNITY ANALYSIS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "8. Keyword Opportunity Analysis", 1)

add_body(doc,
    "Our multi-layer keyword research identified significant untapped search demand across "
    f"{PROSPECT_NAME}'s core service categories. Keywords are categorized by their primary "
    "search layer — SEO, AEO, or GEO — to inform a targeted strategy."
)

add_heading_styled(doc, "Opportunity by Search Layer", 2)
add_chart_image(doc, "layer_distribution.png")

add_heading_styled(doc, "Search Demand by Category", 2)
add_chart_image(doc, "search_demand_by_cluster.png")

add_heading_styled(doc, "Top Keyword Opportunities", 2)
kw_rows = []
# Grab Top 3 SEO Keywords from semrush payload
for kw in market_data.get("prospect", {}).get("top_keywords", [])[:3]:
    kw_rows.append([kw.get("Keyword", ""), "SEO", "Service", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "High"])

# Grab Top 2 AEO Keywords
for kw in market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:2]:
    kw_rows.append([kw.get("Keyword", ""), "AEO", "FAQ", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "Medium"])

# Grab Top 2 GEO Keywords
for kw in market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:2]:
    kw_rows.append([kw.get("Keyword", ""), "GEO", "Informational", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "High"])

if not kw_rows:
    kw_rows = [["Primary Services SEO", "SEO", "General", "10K", "$10", "0.5", "High"]]

add_branded_table(doc,
    headers=["Keyword", "Layer", "Cluster", "Volume", "CPC", "Competition", "Score"],
    rows=kw_rows
)

add_heading_styled(doc, "Opportunity Matrix", 2)
add_chart_image(doc, "opportunity_matrix.png")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 9: INTEGRATED STRATEGY — SEO, AEO & GEO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "9. Integrated Strategy — SEO, AEO & GEO", 1)

add_body(doc,
    "Our recommended strategy addresses all three layers of modern search through "
    "an integrated framework. Each pillar reinforces the others, creating a compounding "
    "effect that maximizes visibility across every search channel."
)

def render_strategy_pillar(doc, title, pillar_data):
    if not pillar_data:
        add_heading_styled(doc, title, 2)
        add_body(doc, "Data unvailable.")
        return
        
    add_heading_styled(doc, title, 2)
    add_body(doc, pillar_data.get("overview", ""))
    
    add_heading_styled(doc, "Key Tactical Initiatives", 3)
    for initiative in pillar_data.get("key_initiatives", []):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(initiative)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GREY

    impact_rows = []
    for item in pillar_data.get("impact_matrix", []):
        impact_rows.append([item.get("opportunity", ""), item.get("expected_outcome", "")])
    
    if impact_rows:
        doc.add_paragraph("")
        add_branded_table(doc, headers=["Opportunity Area", "Expected Business Outcome"], rows=impact_rows)

render_strategy_pillar(doc, "Pillar 1: Technical Foundation & AI Readiness", narrative_data.get("integrated_strategy_technical", {}))
render_strategy_pillar(doc, "Pillar 2: Content & Answer Optimization", narrative_data.get("integrated_strategy_content", {}))
render_strategy_pillar(doc, "Pillar 3: Authority & Entity Building", narrative_data.get("integrated_strategy_authority", {}))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 10: CONTENT STRATEGY
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "10. Content Strategy", 1)

add_body(doc,
    "A well-executed content strategy is the engine that drives growth across all three "
    "search layers. Based on our AI's analysis of your search landscape and competitive gaps, we have formulated customized editorial pillars:"
)

content_roadmap = narrative_data.get("content_strategy_roadmap", [])
if content_roadmap:
    for item in content_roadmap:
        add_heading_styled(doc, item.get("topic", "Content Pillar"), 2)
        add_body(doc, item.get("rationale", ""))
        
        doc.add_paragraph("")
        add_branded_table(doc, 
            headers=["Target Audience Intent", "Recommended Schemas"], 
            rows=[[item.get("target_audience_intent", ""), ", ".join(item.get("recommended_schemas", []))]]
        )
        
        doc.add_paragraph("")
        add_heading_styled(doc, "Recommended Sub-Topics (H2 Structure)", 3)
        for sub in item.get("sub_topics", []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(sub)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
else:
    add_body(doc, "Expand technical and location-based clusters targeting high-intent long-tail keywords.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 11: 90-DAY ACTION PLAN
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "11. 90-Day Action Plan", 1)

add_body(doc,
    "The following phased roadmap outlines the key initiatives for the first 90 days, "
    "with actions organized across all three search layers."
)

add_branded_table(doc,
    headers=["Phase", "SEO Actions", "AEO Actions", "GEO Actions"],
    rows=[
        ["Month 1\nFoundation",
         "Fix critical technical issues, optimize title tags and meta descriptions",
         "Implement FAQPage schema, add FAQ sections to top service pages",
         "Ensure AI bots are not blocked, implement Organization schema, create llms.txt"],
        ["Month 2\nContent",
         "Publish cornerstone content for top clusters, optimize existing pages",
         "Create how-to guides for top question keywords, add Q&A structured content",
         "Publish thought leadership content, strengthen author bios and entity signals"],
        ["Month 3\nAuthority",
         "Launch outreach campaign, amplify top content",
         "Optimize for PAA opportunities, expand FAQ coverage",
         "Build entity consistency across web, pursue digital PR for citations"],
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 12: WHY TRAFFICRADIUS & NEXT STEPS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "12. Why TrafficRadius & Next Steps", 1)

add_heading_styled(doc, "Why TrafficRadius", 2)
add_body(doc,
    "TrafficRadius is a data-driven digital growth agency that specializes in helping "
    "businesses dominate across all three layers of modern search. Our integrated SEO, AEO, "
    "and GEO approach ensures your brand is visible wherever your customers are searching — "
    "whether in traditional results, answer boxes, or AI-generated responses."
)

add_heading_styled(doc, "Next Steps", 2)
add_body(doc,
    "We are excited about the opportunity to partner with " + PROSPECT_NAME + " and help "
    "you achieve multi-channel search dominance. To get started:"
)

add_branded_table(doc,
    headers=["Step", "Action", "Timeline"],
    rows=[
        ["1", "Review this strategy document and the accompanying data workbook", "This week"],
        ["2", "Schedule a strategy call to discuss findings and answer questions", "Next week"],
        ["3", "Approve the engagement and kick off the project", "Week 3"],
    ]
)

# ── Save ─────────────────────────────────────────────────────

os.makedirs(OUTPUT_DIR, exist_ok=True)
output_path = os.path.join(OUTPUT_DIR, "Strategy_Document.docx")
doc.save(output_path)
print(f"DOCX saved to: {output_path}")
